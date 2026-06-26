"""
word_utils.py - Utilitários para manipulação de placeholders em templates Word

Funções:
- extract_placeholders(template_path): Extrai todos os placeholders {{campo}} de um template
- substitute_placeholders(template_path, output_path, substitutions): Substitui placeholders pelos valores
- get_template_metadata(template_name): Retorna metadados do template (placeholders esperados, campos fixos)
- inserir_conteudo(modelo_path, conteudo, output_path, ...): [LEGACY] Compatibilidade com generator.py
"""

import re
import os
from pathlib import Path
from typing import Dict, List, Set, Any, Callable, Optional
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Padrão regex para encontrar placeholders {{campo}}
PLACEHOLDER_PATTERN = r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}'

# Constantes de altura de imagem por modo
# Tradicional (todos os contratos exceto 1565): 10cm
# SP (c0908):                                   10cm
# SP2 (c1565 exclusivo):                         7cm — sobrescrito em word_utils_sp2.py
ALTURA_PADRAO = 10.0  # cm — padrão para Tradicional e SP (c0908)
LARGURA_MAX_3_COL = 15.0  # cm
LARGURA_MAX_2_COL = 16.0  # cm

# Lista de placeholders esperados (de acordo com PLANEJAMENTO_V3.2.md)
EXPECTED_PLACEHOLDERS = {
    'nr_os',                      # Número da Ordem de Serviço
    'data_elaboracao',            # Data de elaboração (auto-gerada = hoje)
    'data_atendimento',           # Data do atendimento/vistoria
    'agencia_codigo',             # Código/Prefixo da agência
    'agencia_nome',               # Nome da agência
    'endereco',                   # Endereço completo
    'responsavel_dependencia'     # Matrícula + Nome do responsável
}


def extract_placeholders(template_path: str) -> Dict[str, List[str]]:
    """
    Extrai todos os placeholders {{campo}} de um template Word.

    Args:
        template_path: Caminho completo do arquivo .docx

    Returns:
        Dict com:
        - 'placeholders': Set de placeholders encontrados (ex: {'nr_os', 'data_elaboracao'})
        - 'locations': Dict mapeando placeholder para lista de localizações (parágrafo/tabela)
        - 'missing': Set de placeholders esperados mas não encontrados
        - 'extra': Set de placeholders inesperados

    Raises:
        FileNotFoundError: Se template_path não existir
        Exception: Se houver erro ao processar o documento
    """
    try:
        template_path = Path(template_path)
        if not template_path.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        doc = Document(template_path)
        found_placeholders = set()
        locations = {}

        # 1. Procurar em parágrafos
        for para_idx, para in enumerate(doc.paragraphs):
            matches = re.findall(PLACEHOLDER_PATTERN, para.text)
            for match in matches:
                found_placeholders.add(match)
                if match not in locations:
                    locations[match] = []
                locations[match].append(f"paragraph_{para_idx}")

        # 2. Procurar em tabelas
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    matches = re.findall(PLACEHOLDER_PATTERN, cell.text)
                    for match in matches:
                        found_placeholders.add(match)
                        if match not in locations:
                            locations[match] = []
                        locations[match].append(f"table_{table_idx}_row_{row_idx}_col_{col_idx}")

        # 3. Calcular missing e extra
        missing = EXPECTED_PLACEHOLDERS - found_placeholders
        extra = found_placeholders - EXPECTED_PLACEHOLDERS

        result = {
            'placeholders': found_placeholders,
            'locations': locations,
            'missing': missing,
            'extra': extra,
            'total_found': len(found_placeholders),
            'total_expected': len(EXPECTED_PLACEHOLDERS)
        }

        logger.info(f"✓ Extraído de {template_path.name}: {len(found_placeholders)}/{len(EXPECTED_PLACEHOLDERS)} placeholders")

        return result

    except Exception as e:
        logger.error(f"✗ Erro ao extrair placeholders de {template_path}: {str(e)}")
        raise


def substitute_placeholders(
    template_path: str,
    output_path: str,
    substitutions: Dict[str, str],
    validate: bool = True
) -> Dict[str, any]:
    """
    Substitui placeholders {{campo}} pelos valores fornecidos.
    """
    try:
        template_path = Path(template_path)
        output_path = Path(output_path)

        if not template_path.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        # 1. Extrair placeholders do template
        extracted = extract_placeholders(str(template_path))
        found_placeholders = extracted['placeholders']

        # 2. Validação
        if validate:
            if extracted['missing']:
                raise ValueError(
                    f"Template tem placeholders faltando: {extracted['missing']} "
                    f"(esperados {len(EXPECTED_PLACEHOLDERS)}, encontrados {len(found_placeholders)})"
                )

        missing_values = found_placeholders - set(substitutions.keys())
        extra_values = set(substitutions.keys()) - found_placeholders

        # 3. Abrir documento
        doc = Document(template_path)
        substitution_count = 0

        # 4. Substituir em parágrafos
        for para in doc.paragraphs:
            for placeholder, value in substitutions.items():
                pattern = r'\{\{' + placeholder + r'\}\}'
                if re.search(pattern, para.text):
                    old_text = para.text
                    new_text = re.sub(pattern, str(value), para.text)

                    if old_text != new_text:
                        for run in para.runs:
                            run.text = ''
                        para.text = new_text
                        substitution_count += 1

        # 5. Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in substitutions.items():
                        pattern = r'\{\{' + placeholder + r'\}\}'
                        if re.search(pattern, cell.text):
                            old_text = cell.text
                            new_text = re.sub(pattern, str(value), cell.text)

                            if old_text != new_text:
                                cell.text = ''
                                cell.text = new_text
                                substitution_count += 1

        # 6. Salvar documento
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        result = {
            'success': True,
            'substitutions_made': substitution_count,
            'output_path': str(output_path),
            'missing_values': missing_values,
            'extra_values': extra_values,
            'template_path': str(template_path),
            'placeholders_found': len(found_placeholders)
        }

        logger.info(
            f"✓ Substituído {template_path.name} → {output_path.name} "
            f"({substitution_count} substituições)"
        )

        return result

    except Exception as e:
        logger.error(f"✗ Erro ao substituir placeholders: {str(e)}")
        raise


def get_template_metadata(template_name: str) -> Dict[str, any]:
    """
    Retorna metadados de um template específico.
    """
    # Extrair código e nome da agência do nome do arquivo
    parts = template_name.replace('.docx', '').split(' - ')

    if len(parts) >= 3:
        code = parts[1]
        agency = parts[2]
    else:
        code = "UNKNOWN"
        agency = "UNKNOWN"

    metadata = {
        'name': template_name,
        'code': code,
        'agency': agency,
        'placeholders': EXPECTED_PLACEHOLDERS,
        'fields': {
            'nr_os': {
                'label': 'Número da Ordem de Serviço',
                'type': 'string',
                'required': True,
                'example': '1753',
                'category': 'dynamic'
            },
            'data_elaboracao': {
                'label': 'Data de Elaboração',
                'type': 'date',
                'required': True,
                'example': '2026-05-03',
                'category': 'dynamic',
                'auto_fill': True
            },
            'data_atendimento': {
                'label': 'Data do Atendimento',
                'type': 'date',
                'required': True,
                'example': '2026-05-01',
                'category': 'dynamic'
            },
            'agencia_codigo': {
                'label': 'Código da Agência',
                'type': 'string',
                'required': True,
                'example': code,
                'category': 'fixed'
            },
            'agencia_nome': {
                'label': 'Nome da Agência',
                'type': 'string',
                'required': True,
                'example': agency,
                'category': 'fixed'
            },
            'endereco': {
                'label': 'Endereço da Dependência',
                'type': 'string',
                'required': True,
                'example': 'Avenida Brasil, 1000 - São Paulo, SP',
                'category': 'fixed'
            },
            'responsavel_dependencia': {
                'label': 'Responsável da Dependência',
                'type': 'string',
                'required': True,
                'example': '123456 - João Silva',
                'category': 'fixed'
            }
        },
        'dynamic_fields': ['nr_os', 'data_elaboracao', 'data_atendimento'],
        'fixed_fields': ['agencia_codigo', 'agencia_nome', 'endereco', 'responsavel_dependencia']
    }

    return metadata


def validate_substitutions(
    template_name: str,
    substitutions: Dict[str, str]
) -> Dict[str, any]:
    """
    Valida se os valores fornecidos são válidos para o template.
    """
    errors = []
    warnings = []
    missing = set()

    metadata = get_template_metadata(template_name)
    fields = metadata['fields']

    # 1. Verificar placeholders obrigatórios
    for field_name, field_config in fields.items():
        if field_config['required']:
            if field_name not in substitutions:
                if field_config.get('auto_fill'):
                    warnings.append(f"Campo '{field_name}' será auto-preenchido")
                else:
                    missing.add(field_name)
                    errors.append(f"Campo obrigatório faltando: '{field_name}'")

    # 2. Validar tipos e formatos
    for field_name, value in substitutions.items():
        if field_name not in fields:
            warnings.append(f"Campo desconhecido fornecido: '{field_name}'")
            continue

        field_config = fields[field_name]
        field_type = field_config['type']

        if field_type == 'date':
            if not re.match(r'^\d{4}-\d{2}-\d{2}$|^\d{2}/\d{2}/\d{4}$', str(value)):
                errors.append(f"Campo '{field_name}' tem formato de data inválido: '{value}'")

        if field_name in ['agencia_codigo', 'agencia_nome', 'endereco']:
            if not value or len(str(value).strip()) == 0:
                errors.append(f"Campo '{field_name}' não pode estar vazio")

    return {
        'valid': len(errors) == 0,
        'errors': errors,
        'warnings': warnings,
        'missing': missing,
        'fields_provided': len(substitutions),
        'fields_expected': len([f for f in fields if fields[f]['required']])
    }


# ============================================================================
# COMPATIBILIDADE COM LEGACY CODE (generator.py, word_utils_sp.py, etc)
# ============================================================================

def analisar_imagem(caminho: str) -> tuple:
    """
    Analisa dimensões de imagem para otimizar layout.
    Compatibilidade com word_utils_sp.py

    Returns: (largura_cm, altura_cm)
    """
    try:
        from PIL import Image
        if not os.path.exists(caminho):
            return (6.0, 6.0)

        with Image.open(caminho) as img:
            w_px, h_px = img.size
            # Conversão aproximada: 1cm = 37.7952755906 pixels
            w_cm = w_px / 37.7952755906
            h_cm = h_px / 37.7952755906
            return (w_cm, h_cm)
    except Exception:
        return (6.0, 6.0)


def otimizar_layout(conteudo: List[Any]) -> List[Any]:
    """
    Otimiza layout de conteúdo para melhor visualização.
    Compatibilidade com word_utils_sp.py

    Args:
        conteudo: Lista de itens (strings, dicts com imagens, etc)

    Returns:
        Lista otimizada mantendo ordem e estrutura
    """
    return conteudo


def aplicar_estilo(run, tamanho_pt: int = 11, negrito: bool = False, italico: bool = False):
    """
    Aplica estilo a um texto (run).
    Compatibilidade com word_utils_sp.py
    """
    run.font.size = Pt(tamanho_pt)
    run.font.bold = negrito
    run.font.italic = italico
    run.font.name = 'Arial'


def substituir_placeholders(doc: Document, meta: Optional[dict] = None, selected_description: Optional[str] = None):
    """
    Substitui placeholders {{campo}} no documento.
    Compatibilidade com word_utils_sp.py

    Args:
        doc: Documento Word já carregado
        meta: Dicionário com metadados/placeholders para substituir
        selected_description: Descrição opcional
    """
    if not meta:
        return

    # Processa parágrafos
    for para in doc.paragraphs:
        for field, value in meta.items():
            pattern = r'\{\{' + field + r'\}\}'
            if re.search(pattern, para.text):
                para.text = re.sub(pattern, str(value), para.text)

    # Processa tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field, value in meta.items():
                    pattern = r'\{\{' + field + r'\}\}'
                    if re.search(pattern, cell.text):
                        cell.text = re.sub(pattern, str(value), cell.text)


def inserir_conteudo(
    modelo_path: str,
    conteudo: List[Any],
    output_path: str,
    logger: Optional[Callable[[str], None]] = None,
    selected_description: Optional[str] = None,
    meta: Optional[dict] = None
) -> int:
    """
    [LEGACY] Insere conteúdo (imagens, títulos, etc) em um documento Word.

    Esta função mantém compatibilidade com o código existente (generator.py).
    Processa array 'conteudo' contendo:
    - strings (títulos com prefixo »)
    - {"imagem": caminho}
    - {"quebra_pagina": True}
    - {"descricao_texto": "..."}
    - etc

    Args:
        modelo_path: Caminho do template .docx
        conteudo: Array de itens a inserir
        output_path: Caminho de saída do .docx gerado
        logger: Função para logs (opcional)
        selected_description: Descrição selecionada
        meta: Dicionário com metadados para substituir placeholders

    Returns:
        Total de imagens inseridas
    """
    def _log(msg: str):
        if logger:
            logger(msg)
        else:
            logger_instance = logging.getLogger(__name__)
            logger_instance.info(msg)

    try:
        _log(f">>> Abrindo template: {modelo_path}")
        doc = Document(modelo_path)

        # Substitui placeholders
        if meta:
            substituir_placeholders(doc, meta, selected_description)

        contador_imagens = 0
        paragrafo_insercao_index = None

        # Localiza o marcador {{start_here}}
        for i, paragrafo in enumerate(doc.paragraphs):
            if "{{start_here}}" in paragrafo.text:
                paragrafo_insercao_index = i
                paragrafo.text = ""  # Remove o marcador
                break

        if paragrafo_insercao_index is None:
            _log("[AVISO] Marca '{{start_here}}' não encontrada no template.")
            _log("Salvando documento em: " + output_path)
            doc.save(output_path)
            return 0

        i = 0
        while i < len(conteudo):
            item = conteudo[i]
            try:
                # Títulos
                if isinstance(item, str):
                    titulo = item.replace("»", "").strip()
                    nivel = item.count("»")
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(titulo)
                    p.style = f'Heading {min(nivel + 1, 3)}'
                    run = p.runs[0]
                    aplicar_estilo(run, 11, True)
                    i += 1
                # Imagens
                elif isinstance(item, dict) and "imagem" in item:
                    grupo = []
                    while i < len(conteudo) and isinstance(conteudo[i], dict) and "imagem" in conteudo[i]:
                        grupo.append(conteudo[i]["imagem"])
                        i += 1
                    sub_i = 0
                    while sub_i < len(grupo):
                        if sub_i + 1 < len(grupo):
                            w1, h1 = analisar_imagem(grupo[sub_i])
                            w2, h2 = analisar_imagem(grupo[sub_i+1])
                            if h1 > w1 and h2 > w2:
                                table = doc.add_table(rows=1, cols=2)
                                table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                for col in range(2):
                                    cell_p = table.cell(0, col).paragraphs[0]
                                    cell_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    cell_p.add_run().add_picture(grupo[sub_i + col], width=Cm(7.0))
                                    contador_imagens += 1
                                doc.paragraphs[paragrafo_insercao_index]._element.addbefore(table._element)
                                sub_i += 2
                                continue
                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        w, h = analisar_imagem(grupo[sub_i])
                        tw = 10.0 if w > h else 7.0
                        p.add_run().add_picture(grupo[sub_i], width=Cm(tw))
                        contador_imagens += 1
                        sub_i += 1
                elif isinstance(item, dict) and "quebra_pagina" in item:
                    doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(3)
                    i += 1
                else:
                    i += 1
            except Exception as e:
                _log(f"[ERRO] Falha ao processar item: {e}")
                i += 1
                continue

        # Salva documento
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        _log(f">>> Documento salvo: {output_path}")
        _log(f">>> Total de imagens inseridas: {contador_imagens}")

        return contador_imagens

    except Exception as e:
        _log(f"[ERRO CRÍTICO] Falha ao gerar relatório: {e}")
        raise
