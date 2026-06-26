import os
import re
import copy
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from word_utils import (
    ALTURA_PADRAO,
    LARGURA_MAX_3_COL,
    LARGURA_MAX_2_COL,
    analisar_imagem,
    otimizar_layout,
    aplicar_estilo,
    substituir_placeholders,
)
from utils_sp import formatar_moeda_texto

def set_cell_background(cell, color="D9D9D9"):
    """Atribui cor de fundo à célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def inserir_conteudo_sp(modelo_path, conteudo, output_path, selected_description=None, meta=None):
    doc = Document(modelo_path)
    substituir_placeholders(doc, meta, selected_description)

    contador_imagens = 0
    paragrafo_insercao_index = None

    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("[AVISO] Marca '{{start_here}}' não encontrada.")
        return 0
    
    # Aplica otimização de imagens (reutiliza a lógica do Tradicional)
    conteudo_otimizado = otimizar_layout(conteudo)

    # Inverte ordem para inserir "antes" na posição
    conteudo_invertido = list(reversed(conteudo_otimizado))

    for item in conteudo_invertido:
        # Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip()
            nivel = item.count("»")
            
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            if nivel == 0:
                p.style = 'Heading 1'
                p.add_run(f"- {titulo}:")
            elif nivel == 1:
                p.style = 'Heading 2'
                p.add_run(f"{titulo}:")
            elif nivel == 2:
                p.style = 'Heading 3'
                p.add_run(f"{titulo}:")
            
        # Textos Padrão (- Detalhes:, - Vista ampla:)
        elif isinstance(item, dict) and "texto_padrao" in item:
            texto = f"- {item['texto_padrao']}:"
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(texto)
            aplicar_estilo(run, 11, True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
        # Imagem da Fachada (Tamanho maior / Destaque)
        elif isinstance(item, dict) and "imagem_fachada" in item:
            from PIL import Image
            imagem_path = item["imagem_fachada"]
            if os.path.exists(imagem_path):
                try:
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                    run = p.add_run()
                    with Image.open(imagem_path) as img:
                        w, h = img.size
                        # Fachada costuma ser capa, deixamos com 12cm de altura ou proporcional a largura da pagina
                        altura_fachada = 12.0
                        prop = altura_fachada / (h / 37.7952755906)
                        w_cm = (w * prop) / 37.7952755906
                        
                        # Se ficar muito larga pro A4, limita a largura a 16cm
                        if w_cm > 16.0:
                            w_cm = 16.0
                            prop_w = 16.0 / (w / 37.7952755906)
                            altura_fachada = (h * prop_w) / 37.7952755906
                            
                    run.add_picture(imagem_path, width=Cm(w_cm), height=Cm(altura_fachada))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Título de Fachada (Abaixo da foto)
                    p_titulo = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                    run_t = p_titulo.add_run("FACHADA")
                    aplicar_estilo(run_t, 12, True)
                    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    contador_imagens += 1
                except Exception as e:
                    print(f"Erro imagem fachada {imagem_path}: {e}")
            
        # Descrição Técnica Textual (auto-gerada do SP)
        elif isinstance(item, dict) and "texto_descricao" in item:
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            texto_raw = item["texto_descricao"]

            # Processa o marcador <RED>
            partes = re.split(r'(<RED>.*?</RED>)', texto_raw)
            for parte in partes:
                if parte.startswith("<RED>") and parte.endswith("</RED>"):
                    conteudo_red = parte[5:-6]
                    run = p.add_run(conteudo_red)
                    run.font.color.rgb = RGBColor(255, 0, 0) # VERMELHO
                    aplicar_estilo(run, 11, True) # Negrito
                else:
                    run = p.add_run(parte)
                    aplicar_estilo(run, 11, False)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Observação técnica do modal (texto livre do usuário)
        elif isinstance(item, dict) and "descricao_texto" in item:
            texto = item.get("descricao_texto", "").strip()
            if texto:
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                run = p.add_run(texto)
                aplicar_estilo(run, 11, False)  # 11pt, not bold
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_before = Pt(2)  # 40 twips
                p.paragraph_format.space_after = Pt(2)

        # Tabela de Medição Otimizada
        elif isinstance(item, dict) and "tabela_medicao" in item:
            # Pular uma linha antes
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            
            tabela_dados = item["tabela_medicao"]
            
            if tabela_dados["tipo"] == "pintura":
                tabela = doc.add_table(rows=2, cols=5)
                doc.paragraphs[paragrafo_insercao_index]._p.addnext(tabela._tbl)
                
                tabela.style = 'Table Grid'
                tabela.autofit = False
                tabela.allow_autofit = False

                larguras = [3.5, 3.0, 3.0, 3.0, 3.5]
                for i, col in enumerate(tabela.columns):
                    col.width = Cm(larguras[i])

                # Linha de Título Geral (Mesclada)
                titulo_cel = tabela.cell(0, 0)
                for i in range(1, 4):
                    titulo_cel = titulo_cel.merge(tabela.cell(0, i))
                titulo_cel = titulo_cel.merge(tabela.cell(0, 4))
                
                titulo_cel.text = "Pintura em látex acrílica standard fosca sem emassamento, 3 demãos, com aplicação de selador para exterior (ITEM 17.4)"
                titulo_cel.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run_t = titulo_cel.paragraphs[0].runs[0]
                run_t.font.bold = True
                run_t.font.size = Pt(10)
                set_cell_background(titulo_cel, "808080")

                # Linha de Cabeçalho
                hdr_cells = tabela.rows[1].cells
                headers = ["REFERÊNCIA", "LARGURA (m)", "ALTURA (m)", "DESCONTO", "TOTAL (m²)"]
                for i, text in enumerate(headers):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)
                    set_cell_background(hdr_cells[i], "D9D9D9")

                soma_bruta = 0.0
                soma_desconto = 0.0

                for med in tabela_dados["medidas"]:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = med["referencia"]
                    row_cells[1].text = formatar_moeda_texto(med["largura"])
                    row_cells[2].text = formatar_moeda_texto(med["altura"])
                    row_cells[3].text = formatar_moeda_texto(med["desconto"])
                    row_cells[4].text = formatar_moeda_texto(med["subtotal"])
                    
                    for c in row_cells:
                        c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    soma_bruta += med["subtotal"]
                    soma_desconto += med["desconto"]

                # Linha Subtotal
                row_sub = tabela.add_row().cells
                sub_label = row_sub[0].merge(row_sub[1]).merge(row_sub[2])
                sub_label.text = "Subtotal"
                sub_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                sub_label.paragraphs[0].runs[0].font.bold = True
                set_cell_background(sub_label, "D9D9D9")
                
                row_sub[3].text = formatar_moeda_texto(soma_desconto)
                row_sub[4].text = formatar_moeda_texto(soma_bruta)
                for i in [3, 4]:
                    row_sub[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    set_cell_background(row_sub[i], "D9D9D9")

                # Linha Total
                row_total = tabela.add_row().cells
                total_label = row_total[0].merge(row_total[1]).merge(row_total[2]).merge(row_total[3])
                total_label.text = "Total"
                total_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                total_label.paragraphs[0].runs[0].font.bold = True
                set_cell_background(total_label, "D9D9D9")

                area_liquida = soma_bruta - soma_desconto
                row_total[4].text = formatar_moeda_texto(area_liquida)
                row_total[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row_total[4].paragraphs[0].runs[0].font.bold = True
                set_cell_background(row_total[4], "D9D9D9")

            elif tabela_dados["tipo"] == "mobiliario":
                tabela = doc.add_table(rows=2, cols=2)
                doc.paragraphs[paragrafo_insercao_index]._p.addnext(tabela._tbl)
                tabela.style = 'Table Grid'
                tabela.autofit = False
                tabela.width = Cm(16.0)
                
                # Titulo Mobiliario
                tit_mob = tabela.cell(0, 0).merge(tabela.cell(0, 1))
                tit_mob.text = "Deslocamento ou remanejamento de mobiliário dentro da agência (ITEM 13.12)"
                tit_mob.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                tit_mob.paragraphs[0].runs[0].font.bold = True
                set_cell_background(tit_mob, "808080")

                # Header
                h_cells = tabela.rows[1].cells
                h_cells[0].text = "REFERÊNCIA"
                h_cells[1].text = "TOTAL (UN)"
                for c in h_cells: 
                    c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    c.paragraphs[0].runs[0].font.bold = True
                    set_cell_background(c, "D9D9D9")

                soma_un = 0.0
                for med in tabela_dados["medidas"]:
                    r_cells = tabela.add_row().cells
                    r_cells[0].text = med["referencia"]
                    r_cells[1].text = formatar_moeda_texto(med["total_un"])
                    for c in r_cells: c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    soma_un += med["total_un"]

                r_tot = tabela.add_row().cells
                r_tot[0].text = "Total"
                r_tot[1].text = formatar_moeda_texto(soma_un)
                for c in r_tot: 
                    c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    c.paragraphs[0].runs[0].font.bold = True
                    set_cell_background(c, "D9D9D9")

            # Espaçamento extra pós tabela
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')

        # Imagens Simples
        elif isinstance(item, dict) and "imagem" in item:
            from PIL import Image
            imagem_path = item["imagem"]
            if os.path.exists(imagem_path):
                try:
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                    run = p.add_run()
                    with Image.open(imagem_path) as img:
                        w, h = img.size
                        prop = ALTURA_PADRAO / (h / 37.7952755906)
                        w_cm = (w * prop) / 37.7952755906
                    run.add_picture(imagem_path, width=Cm(w_cm), height=Cm(ALTURA_PADRAO))
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    contador_imagens += 1
                except Exception as e:
                    print(f"Erro imagem {imagem_path}: {e}")
                    
        # Tabela Imagens (Reaproveitado do util base)
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            tabela = doc.add_table(rows=1, cols=colunas)
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)
            
            w_col = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(w_col)
                
            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].add_run()
                    try:
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                    except Exception:
                        pass
        
        # Quebra de Página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

    # Remover a marcação final para não sujar o relatório
    doc.paragraphs[paragrafo_insercao_index].text = doc.paragraphs[paragrafo_insercao_index].text.replace("{{start_here}}", "")
    
    doc.save(output_path)
    return contador_imagens
