# word_utils_sp2.py — AutoRelatório V3.2 · Modo São Paulo 2
# Motor de geração Word para o padrão do Contrato 1565 (São José do Rio Preto)
#
# DIFERENÇAS vs word_utils_sp.py (SP1):
#   ✦ Cabeçalho institucional em TODAS as páginas (seção de cabeçalho do docx)
#   ✦ Tabela de Memória de Cálculo com múltiplas linhas (Parede 1, Parede 2...)
#   ✦ Suporte a faces (coluna FACES na tabela quando > 1)
#   ✦ Enunciado textual (Item 17.X – desc...) antes de cada tabela
#   ✦ Tabela de Itens separada (código | descrição | qtd | unidade)
#   ✦ Inserção de Croquis com legenda centralizada
#   ✦ Rodapé com paginação (CONTRATO X — Empresa / Pág N de Total)
#
# ITENS DO array conteudo processados aqui:
#   str                       → título (Heading 1/2/3 por contagem de »)
#   imagem_fachada            → foto de capa centralizada
#   imagem                    → foto normal
#   croqui                    → imagem de croqui + legenda abaixo
#   texto_padrao              → "- Detalhes:" em negrito
#   texto_narrativo           → parágrafo de texto livre
#   enunciado_item            → "Item 17.X – descrição..." em itálico
#   memoria_calculo           → tabela de cálculo com linhas múltiplas
#   tabela_itens_sp2          → tabela de itens separada
#   descricao_texto           → texto do modal (legado SP1 compatível)
#   tabela_imagens            → grade de imagens verticais (herdado do SP1)
#   quebra_pagina             → quebra de página

import os
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

# Reutiliza helpers do word_utils base
from word_utils import (
    LARGURA_MAX_3_COL,
    LARGURA_MAX_2_COL,
    analisar_imagem,
    otimizar_layout,
    aplicar_estilo,
    substituir_placeholders,
)

# SP2 (c1565) usa altura de imagem menor que os demais modos
# Tradicional e SP usam 10cm; SP2 usa 7cm
ALTURA_PADRAO = 7.0  # cm — exclusivo do contrato 1565
from utils_sp2 import formatar_moeda

# ---------------------------------------------------------------------------
# CORES (padrão Contrato 1565)
# ---------------------------------------------------------------------------

COR_CABECALHO_TABELA = "595959"   # cinza escuro para título da tabela
COR_HEADER_COLUNAS   = "D9D9D9"   # cinza claro para cabeçalho de colunas
COR_TOTAL_ROW        = "D9D9D9"   # cinza claro para linha de total
COR_ITENS_HEADER     = "808080"   # cinza médio para header da tabela de itens


# ---------------------------------------------------------------------------
# HELPERS DE CÉLULA
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: int = 10,
                   align=WD_PARAGRAPH_ALIGNMENT.CENTER, italic: bool = False):
    para = cell.paragraphs[0]
    para.alignment = align
    para.clear()
    run = para.add_run(str(text))
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.name  = 'Arial'


def _set_col_width(tabela, col_widths_cm: list):
    for i, w in enumerate(col_widths_cm):
        for row in tabela.rows:
            row.cells[i].width = Cm(w)


# ---------------------------------------------------------------------------
# INSERÇÃO DE IMAGEM SEGURA
# ---------------------------------------------------------------------------

def _inserir_imagem(doc, paragrafo_ref, caminho: str,
                    altura_cm: float = ALTURA_PADRAO,
                    centralizar: bool = True) -> bool:
    """Insere imagem antes de paragrafo_ref. Retorna True se OK."""
    if not os.path.exists(caminho):
        return False
    try:
        p = paragrafo_ref.insert_paragraph_before('')
        run = p.add_run()
        with Image.open(caminho) as img:
            w, h = img.size
            prop  = altura_cm / (h / 37.7952755906)
            w_cm  = (w * prop) / 37.7952755906
        run.add_picture(caminho, width=Cm(w_cm), height=Cm(altura_cm))
        if centralizar:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return True
    except Exception as e:
        print(f"[ERRO imagem SP2] {caminho}: {e}")
        return False


# ---------------------------------------------------------------------------
# TABELA DE MEMÓRIA DE CÁLCULO
# ---------------------------------------------------------------------------

def _inserir_memoria_calculo(doc, paragrafo_ref, dados: dict):
    """
    Insere a tabela de Memória de Cálculo no padrão Contrato 1565.

    Colunas para tipo 'area':
      FOTO | COMPRIMENTO (m) | ALTURA (m) | DESCONTO (m²) | TOTAL (m²)

    Colunas para tipo 'metalico' (com faces):
      FOTO | COMPRIMENTO (m) | ALTURA (m) | FACES | TOTAL (m²)

    Colunas para tipo 'unitario':
      FOTO | QUANTIDADE | TOTAL

    Sempre: linha de Subtotal (se > 1 linha) + linha de Total.
    """
    linhas     = dados.get("linhas", [])
    tipo       = dados.get("tipo_item", "area")
    descricao  = dados.get("descricao", "")
    codigo     = dados.get("codigo", "")
    total_geral = dados.get("total_geral", 0.0)

    if not linhas:
        return

    # Cabeçalho da tabela (título cinza escuro)
    titulo_tabela = descricao if descricao else codigo

    # Define colunas conforme tipo
    if tipo == "metalico":
        headers    = ["FOTO", "COMPRIMENTO (m)", "ALTURA (m)", "FACES", "TOTAL (m²)"]
        col_widths = [3.5, 3.0, 3.0, 2.0, 2.5]
    elif tipo == "unitario":
        headers    = ["FOTO", "QUANTIDADE", "TOTAL"]
        col_widths = [5.0, 5.0, 4.0]
    else:  # area (padrão)
        headers    = ["FOTO", "COMPRIMENTO (m)", "ALTURA (m)", "DESCONTO (m²)", "TOTAL (m²)"]
        col_widths = [3.5, 3.0, 3.0, 2.5, 2.0]

    n_cols = len(headers)
    n_data = len(linhas)
    # Linhas: 1 título + 1 cabeçalho + dados + (subtotal se >1) + total
    n_subtotal = 1 if n_data > 1 else 0
    n_rows = 2 + n_data + n_subtotal + 1

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    paragrafo_ref._p.addnext(tbl._tbl)
    tbl.style    = 'Table Grid'
    tbl.autofit  = False
    _set_col_width(tbl, col_widths)

    row_idx = 0

    # ── Linha 0: título mesclado (cinza escuro) ──────────────────────────
    titulo_cell = tbl.rows[row_idx].cells[0]
    for c in range(1, n_cols):
        titulo_cell = titulo_cell.merge(tbl.rows[row_idx].cells[c])
    titulo_cell.text = titulo_tabela
    titulo_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = titulo_cell.paragraphs[0].runs[0]
    r.font.bold  = True
    r.font.size  = Pt(10)
    r.font.name  = 'Arial'
    _set_cell_bg(titulo_cell, COR_CABECALHO_TABELA)
    row_idx += 1

    # ── Linha 1: cabeçalho de colunas (cinza claro) ─────────────────────
    for c, h in enumerate(headers):
        cell = tbl.rows[row_idx].cells[c]
        _set_cell_bg(cell, COR_HEADER_COLUNAS)
        _set_cell_text(cell, h, bold=True, size_pt=10)
    row_idx += 1

    # ── Linhas de dados ──────────────────────────────────────────────────
    soma_subtotal = 0.0
    soma_desconto = 0.0

    for linha in linhas:
        cells = tbl.rows[row_idx].cells
        if tipo == "metalico":
            _set_cell_text(cells[0], linha.get("referencia", "Foto"))
            _set_cell_text(cells[1], formatar_moeda(linha.get("largura", 0.0)))
            _set_cell_text(cells[2], formatar_moeda(linha.get("altura",  0.0)))
            _set_cell_text(cells[3], str(int(linha.get("faces", 1))))
            _set_cell_text(cells[4], formatar_moeda(linha.get("total",   0.0)), bold=True)
        elif tipo == "unitario":
            _set_cell_text(cells[0], linha.get("referencia", "Foto"))
            _set_cell_text(cells[1], formatar_moeda(linha.get("quantidade", 1.0)))
            _set_cell_text(cells[2], formatar_moeda(linha.get("total",   0.0)), bold=True)
        else:  # area
            _set_cell_text(cells[0], linha.get("referencia", "Foto"))
            _set_cell_text(cells[1], formatar_moeda(linha.get("largura", 0.0)))
            _set_cell_text(cells[2], formatar_moeda(linha.get("altura",  0.0)))
            _set_cell_text(cells[3], formatar_moeda(linha.get("desconto", 0.0)))
            _set_cell_text(cells[4], formatar_moeda(linha.get("subtotal", 0.0)))

        soma_subtotal += linha.get("subtotal", linha.get("total", 0.0))
        soma_desconto += linha.get("desconto", 0.0)
        row_idx += 1

    # ── Linha de Subtotal (somente se múltiplas linhas e tipo area) ──────
    if n_data > 1 and tipo == "area":
        cells = tbl.rows[row_idx].cells
        sub_label = cells[0].merge(cells[1]).merge(cells[2])
        sub_label.text = "Subtotal"
        sub_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_label.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(sub_label, COR_TOTAL_ROW)
        _set_cell_text(cells[3], formatar_moeda(soma_desconto), bold=True)
        _set_cell_text(cells[4], formatar_moeda(soma_subtotal), bold=True)
        _set_cell_bg(cells[3], COR_TOTAL_ROW)
        _set_cell_bg(cells[4], COR_TOTAL_ROW)
        row_idx += 1

    # ── Linha de Total ───────────────────────────────────────────────────
    cells = tbl.rows[row_idx].cells
    if tipo == "area":
        total_label = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
        total_label.text = "Total"
        total_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        total_label.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(total_label, COR_TOTAL_ROW)
        area_liquida = round(soma_subtotal - soma_desconto, 2)
        _set_cell_text(cells[4], formatar_moeda(area_liquida), bold=True)
        _set_cell_bg(cells[4], COR_TOTAL_ROW)
    else:
        total_label = cells[0].merge(cells[1]) if tipo == "unitario" else cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
        total_label.text = "Total"
        total_label.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        total_label.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(total_label, COR_TOTAL_ROW)
        last_col = cells[2] if tipo == "unitario" else cells[4]
        _set_cell_text(last_col, formatar_moeda(total_geral), bold=True)
        _set_cell_bg(last_col, COR_TOTAL_ROW)


# ---------------------------------------------------------------------------
# TABELA DE ITENS SP2
# ---------------------------------------------------------------------------

def _inserir_tabela_itens_sp2(doc, paragrafo_ref, dados: dict):
    """
    Insere a tabela de Itens separada (padrão Contrato 1565).
    Colunas: Itens | Itens | Itens | Itens  (mescladas no header)
    Linha de dados: código | descrição completa | quantidade | unidade
    """
    codigo    = dados.get("codigo",    "—")
    descricao = dados.get("descricao", "")
    qtd       = dados.get("quantidade", 0.0)
    unidade   = dados.get("unidade",    "m²")

    tbl = doc.add_table(rows=2, cols=4)
    paragrafo_ref._p.addnext(tbl._tbl)
    tbl.style   = 'Table Grid'
    tbl.autofit = False
    _set_col_width(tbl, [2.0, 9.0, 2.5, 2.5])

    # Linha 0: cabeçalho "Itens" mesclado
    header_cell = tbl.rows[0].cells[0]
    for c in range(1, 4):
        header_cell = header_cell.merge(tbl.rows[0].cells[c])
    header_cell.text = "Itens"
    header_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = header_cell.paragraphs[0].runs[0]
    r.font.bold = True
    r.font.size = Pt(10)
    _set_cell_bg(header_cell, COR_ITENS_HEADER)

    # Linha 1: dados
    data = tbl.rows[1].cells
    _set_cell_text(data[0], codigo,    bold=True, size_pt=10)
    _set_cell_text(data[1], descricao, size_pt=10, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    _set_cell_text(data[2], formatar_moeda(qtd) if isinstance(qtd, float) else str(qtd), bold=True, size_pt=10)
    _set_cell_text(data[3], unidade,   size_pt=10)


# ---------------------------------------------------------------------------
# MOTOR PRINCIPAL DE INSERÇÃO
# ---------------------------------------------------------------------------

def inserir_conteudo_sp2(
    modelo_path: str,
    conteudo: list,
    output_path: str,
    meta: dict = None,
    selected_description: str = None
) -> int:
    """
    Insere todo o conteudo do SP2 no template Word e salva em output_path.
    Retorna o total de imagens inseridas.

    meta: dict com dados do cabeçalho institucional (ver generator_sp2.py).
    """
    doc = Document(modelo_path)
    substituir_placeholders(doc, meta, selected_description)

    # ── Localiza marca de inserção ────────────────────────────────────────
    idx_insercao = None
    for i, p in enumerate(doc.paragraphs):
        if "{{start_here}}" in p.text:
            idx_insercao = i
            break

    if idx_insercao is None:
        print("[AVISO SP2] Marca '{{start_here}}' não encontrada no template.")
        return 0

    # ── Otimização de layout (agrupa verticais) ───────────────────────────
    conteudo_otimizado  = otimizar_layout(conteudo)
    conteudo_invertido  = list(reversed(conteudo_otimizado))

    contador_imagens = 0

    for item in conteudo_invertido:
        p_ref = doc.paragraphs[idx_insercao]

        # ── Título de seção ──────────────────────────────────────────────
        if isinstance(item, str):
            titulo = item.replace("»", "").strip()
            nivel  = item.count("»")
            p = p_ref.insert_paragraph_before('')
            if nivel == 0:
                p.style = 'Heading 1'
                p.add_run(f"{titulo}:")
            elif nivel == 1:
                p.style = 'Heading 2'
                p.add_run(f"{titulo}:")
            else:
                p.style = 'Heading 3'
                p.add_run(f"{titulo}:")

        # ── Imagem de Fachada ────────────────────────────────────────────
        elif isinstance(item, dict) and "imagem_fachada" in item:
            caminho = item["imagem_fachada"]
            if _inserir_imagem(doc, p_ref, caminho, altura_cm=12.0):
                contador_imagens += 1
                # Título "FACHADA" abaixo da foto
                p_titulo = p_ref.insert_paragraph_before('')
                run_t = p_titulo.add_run("FACHADA")
                aplicar_estilo(run_t, 12, True)
                p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # ── Foto normal ──────────────────────────────────────────────────
        elif isinstance(item, dict) and "imagem" in item:
            if _inserir_imagem(doc, p_ref, item["imagem"]):
                contador_imagens += 1

        # ── Croqui ───────────────────────────────────────────────────────
        elif isinstance(item, dict) and "croqui" in item:
            caminho = item["croqui"]
            legenda = item.get("legenda", "CROQUI")
            if _inserir_imagem(doc, p_ref, caminho, altura_cm=8.0):
                contador_imagens += 1
            # Legenda abaixo do croqui
            p_leg = p_ref.insert_paragraph_before('')
            run_l = p_leg.add_run(legenda)
            aplicar_estilo(run_l, 9, True)
            p_leg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # ── Texto padrão (- Detalhes:) ───────────────────────────────────
        elif isinstance(item, dict) and "texto_padrao" in item:
            p = p_ref.insert_paragraph_before('')
            run = p.add_run(f"- {item['texto_padrao']}:")
            aplicar_estilo(run, 11, True)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # ── Texto narrativo (descrição do problema) ───────────────────────
        elif isinstance(item, dict) and "texto_narrativo" in item:
            texto = item["texto_narrativo"].strip()
            if texto:
                p = p_ref.insert_paragraph_before('')
                run = p.add_run(texto)
                aplicar_estilo(run, 11, False)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # ── Texto do modal (legado SP1 compatível) ────────────────────────
        elif isinstance(item, dict) and "descricao_texto" in item:
            texto = item["descricao_texto"].strip()
            if texto:
                p = p_ref.insert_paragraph_before('')
                partes = re.split(r'(<RED>.*?</RED>)', texto)
                for parte in partes:
                    if parte.startswith("<RED>") and parte.endswith("</RED>"):
                        run = p.add_run(parte[5:-6])
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        aplicar_estilo(run, 11, True)
                    else:
                        run = p.add_run(parte)
                        aplicar_estilo(run, 11, False)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # ── Enunciado do item (ex: "Item 17.6 – Pintura...") ─────────────
        elif isinstance(item, dict) and "enunciado_item" in item:
            dados = item["enunciado_item"]
            texto = f"Item {dados.get('codigo', '')} – {dados.get('descricao', '')}."
            p = p_ref.insert_paragraph_before('')
            run = p.add_run(texto)
            aplicar_estilo(run, 11, False)
            run.font.italic = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)

        # ── Memória de Cálculo (tabela principal SP2) ─────────────────────
        elif isinstance(item, dict) and "memoria_calculo" in item:
            p_ref.insert_paragraph_before('')  # espaçamento antes
            _inserir_memoria_calculo(doc, p_ref, item["memoria_calculo"])
            p_ref.insert_paragraph_before('')  # espaçamento depois

        # ── Tabela de Itens separada ──────────────────────────────────────
        elif isinstance(item, dict) and "tabela_itens_sp2" in item:
            p_ref.insert_paragraph_before('')
            _inserir_tabela_itens_sp2(doc, p_ref, item["tabela_itens_sp2"])
            p_ref.insert_paragraph_before('')

        # ── Tabela de imagens agrupadas (verticais) ───────────────────────
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]
            p_sep   = p_ref.insert_paragraph_before('')
            tabela  = doc.add_table(rows=1, cols=colunas)
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)
            w_col = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(w_col)
            for idx_i, img_path in enumerate(imagens):
                if idx_i < len(tabela.columns):
                    cell = tabela.cell(0, idx_i)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].add_run()
                    try:
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        contador_imagens += 1
                    except Exception:
                        pass

        # ── Quebra de página ──────────────────────────────────────────────
        elif isinstance(item, dict) and "quebra_pagina" in item:
            p_ref.insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

    # Limpa a marca de inserção
    doc.paragraphs[idx_insercao].text = (
        doc.paragraphs[idx_insercao].text.replace("{{start_here}}", "")
    )

    doc.save(output_path)
    print(f"[SP2] Relatório salvo: {output_path} ({contador_imagens} imagens)")
    return contador_imagens
